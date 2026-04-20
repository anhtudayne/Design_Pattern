#!/usr/bin/env python3
"""Sinh hai file Word báo cáo (bố cục giống docs/patterns/08-dynamic-pricing-engine-report-bo-cuc-mau.md)."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

def _add_quote_block(doc, lines):
    try:
        p = doc.add_paragraph(style="Intense Quote")
    except KeyError:
        p = doc.add_paragraph()
    for i, (text, bold_label) in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        if bold_label and ":" in text:
            label, rest = text.split(":", 1)
            r = p.add_run(label + ":")
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(text)

def _add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.size = Pt(11)

def _add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.add_run("—" * 20)

def build_composite_doc():
    doc = Document()
    doc.add_heading("Composite (Dashboard thống kê) — Trình bày theo bố cục mẫu", level=1)
    _add_quote_block(doc, [
        ("Luồng: GET /api/admin/dashboard/stats — gom số liệu tổng quan cho trang quản trị", True),
        ("Tài liệu kỹ thuật chi tiết: docs/patterns/composite-dashboard-stats-package-vi.md và docs/patterns/05-composite.md", True),
    ])
    _add_separator(doc)
    doc.add_heading("Các pattern áp dụng", level=2)
    _add_body(doc, "Trong hệ thống đặt vé rạp chiếu phim, nhóm xử lý thống kê dashboard admin áp dụng Composite Pattern: một interface chung StatsComponent (vai trò Component) cho phép mọi nút — dù là một chỉ số đơn (Leaf) hay bộ gom (Composite) — đều thực hiện cùng một thao tác collect(Map) để ghi dữ liệu vào một map phản hồi. Các Leaf cụ thể (MovieStatsLeaf, UserStatsLeaf, ShowtimeStatsLeaf, FnbStatsLeaf, TicketStatsLeaf, PromotionStatsLeaf, RevenueStatsLeaf) mỗi class chỉ lo một loại số liệu (đếm repository hoặc cộng doanh thu). DashboardStatsComposite nhận toàn bộ bean StatsComponent do Spring inject, lọc bỏ chính nó để tránh đệ quy, rồi gọi lần lượt collect trên cùng map đích. Controller chỉ cần một dòng gọi composite.")
    _add_separator(doc)
    doc.add_heading("Lý do sử dụng", level=2)
    _add_body(doc, "Trang admin cần đồng thời nhiều chỉ số (phim, user, suất chiếu, F&B, vé, khuyến mãi, doanh thu). Nếu viết tất cả lời gọi repository vào một controller hoặc một method dài, mã sẽ khó đọc, khó test, và mỗi lần thêm chỉ số mới phải mở lại lớp điều phối — trái với hướng mở rộng an toàn.")
    _add_body(doc, "Composite giải quyết việc client (ở đây là DashboardController) chỉ muốn một điểm gọi nhưng bên trong có nhiều nguồn dữ liệu: controller tạo HashMap và gọi dashboardStatsComposite.collect(stats) một lần; bên trong, từng Leaf tự put key riêng (totalMovies, totalTickets, …) nên không trùng trách nhiệm.")
    _add_body(doc, "Spring đóng vai trò dựng sẵn danh sách các StatsComponent: mỗi Leaf là @Component; composite nhận List<StatsComponent> ở constructor — thêm Leaf mới không cần sửa danh sách thủ công trong composite.")
    _add_separator(doc)
    doc.add_heading("Ưu điểm", level=2)
    _add_body(doc, "Cấu trúc đóng với thay đổi ở controller, mở với mở rộng chỉ số: thêm một file Leaf mới implement StatsComponent là dashboard có thêm số liệu mà không bắt buộc chỉnh DashboardStatsComposite hay controller.")
    _add_body(doc, "Mã dễ kiểm thử hơn: có thể test từng Leaf với repository giả lập; có thể test composite với danh sách StatsComponent giả. Hợp đồng API thống nhất — một JSON map — trong khi bên trong tách mắt xích rõ vai trò.")
    _add_body(doc, "Composite còn giúp tránh lỗi đệ quy nhờ bước lọc instanceof DashboardStatsComposite khi Spring vô tình đưa cả bean composite vào list (vì nó cũng implement StatsComponent). Tổng thể phù hợp bảo trì và mở rộng dashboard lâu dài.")
    return doc

def build_singleton_doc():
    doc = Document()
    doc.add_heading("Singleton (RestTemplate qua Spring IoC) — Trình bày theo bố cục mẫu", level=1)
    _add_quote_block(doc, [
        ("Luồng: Khởi động ứng dụng Spring — tạo một bean RestTemplate (singleton scope mặc định); MomoServiceImpl dùng cùng instance khi POST tới cổng MoMo (trừ khi bật dev-skip-external)", True),
        ("Tài liệu kỹ thuật chi tiết: docs/patterns/singleton-resttemplate-package-vi.md và docs/patterns/06-singleton.md", True),
    ])
    _add_separator(doc)
    doc.add_heading("Các pattern áp dụng", level=2)
    _add_body(doc, "Trong backend thanh toán MoMo, dự án áp dụng Singleton thông qua Spring IoC: class RestTemplateConfig khai báo @Bean method restTemplate() trả về một RestTemplate mà Spring container giữ trong singleton scope (mặc định). Service MomoServiceImpl nhận RestTemplate qua constructor injection — không gọi new RestTemplate() rải rác trong từng request. Đây là biến thể Singleton do container quản lý, khác Singleton cổ điển (enum/getInstance trong chính class).")
    _add_separator(doc)
    doc.add_heading("Lý do sử dụng", level=2)
    _add_body(doc, "Gọi HTTP ra cổng đối tác (MoMo) cần client ổn định. Tạo new RestTemplate() mỗi lần gọi gây overhead (tạo object, pool kết nối), khó cấu hình thống nhất (timeout, TLS, interceptor) và dễ không nhất quán giữa các lớp.")
    _add_body(doc, "Spring Singleton bean gom việc tạo instance một lần khi context khởi động; mọi service cần HTTP có thể inject cùng bean — phù hợp tài nguyên dùng chung và một điểm mở rộng cấu hình trong RestTemplateConfig.")
    _add_separator(doc)
    doc.add_heading("Ưu điểm", level=2)
    _add_body(doc, "Tiết kiệm tài nguyên và thống nhất cấu hình: một instance dùng chung, có thể bổ sung timeout hoặc factory tại một chỗ (@Bean method) mà không sửa nhiều service.")
    _add_body(doc, "Dễ bảo trì và mở rộng: service mới chỉ cần khai báo phụ thuộc RestTemplate trong constructor; tránh anti-pattern new trong method. Phù hợp vận hành lâu dài khi số tích hợp HTTP tăng.")
    _add_body(doc, "Luồng nghiệp vụ vẫn rõ ràng: MomoServiceImpl tập trung logic ký HMAC và payload MoMo; phần vận chuyển HTTP dùng client đã inject. Có thể bật momo.dev-skip-external để dev không gọi ngoài mà vẫn giữ cùng wiring bean.")
    return doc

def main():
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docs" / "patterns"
    out_dir.mkdir(parents=True, exist_ok=True)
    build_composite_doc().save(out_dir / "composite-dashboard-stats-report-bo-cuc-mau.docx")
    build_singleton_doc().save(out_dir / "singleton-resttemplate-report-bo-cuc-mau.docx")
    print("OK:", out_dir / "composite-dashboard-stats-report-bo-cuc-mau.docx")
    print("OK:", out_dir / "singleton-resttemplate-report-bo-cuc-mau.docx")

if __name__ == "__main__":
    main()
